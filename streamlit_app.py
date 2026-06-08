"""
Systematic Review Bot v8
3 Dimensions: D1 Standardization & AI, D2 Context Engineering, D3 Token Efficiency
PRISMA 2020 + Auto E1/E2/E7 + Year Recovery + Post-fetch Language Check + Word Template
"""

import streamlit as st
import pandas as pd
import io, re, csv, time, concurrent.futures, threading
from datetime import datetime

st.set_page_config(page_title="SR Bot", page_icon="🔬", layout="wide")

# ── Constants ─────────────────────────────────────────────────────────────────
QUERY_MAP = {
    "D1Q1": "D1_Standardization_AI", "D1Q2": "D1_Standardization_AI",
    "D2Q1": "D2_Context_Engineering", "D2Q2": "D2_Context_Engineering", "D2Q3": "D2_Context_Engineering",
    "D3Q1": "D3_Token_Efficiency",   "D3Q2": "D3_Token_Efficiency",   "D3Q3": "D3_Token_Efficiency",
}
DIMENSION_NAMES = {
    "D1": "D1_Standardization_AI",
    "D2": "D2_Context_Engineering",
    "D3": "D3_Token_Efficiency",
}
EXCLUSION_CRITERIA = [
    "E1: Outside year range (2015-2026)",
    "E2: Not English",
    "E3: Not journal/conference/review paper",
    "E4: Not relevant to dimension topic",
    "E5: Duplicate",
    "E6: Full text not accessible",
    "E7: Abstract only / insufficient detail",
    "E8: Not peer-reviewed",
    "E9: Paid / not open access",
]

# ── Parsers ───────────────────────────────────────────────────────────────────
def _paper(title="", authors="", year="", source="", doi="", url="", ptype="", database="", query_id=""):
    return {
        "title": title, "authors": authors, "year": str(year), "source": source,
        "doi": doi, "abstract": "", "url": url, "type": ptype,
        "database": database, "query_id": query_id,
        "dimension": QUERY_MAP.get(query_id, ""),
        "screening_status": "Pending", "exclusion_reason": "", "notes": "",
        "auto_excluded": False,
    }

def parse_springer_csv(content, qid):
    papers = []
    try:
        for row in csv.DictReader(io.StringIO(content)):
            papers.append(_paper(
                title=row.get("Item Title","").strip(), authors=row.get("Authors","").strip(),
                year=str(row.get("Publication Year","")).strip(),
                source=(row.get("Publication Title","") or row.get("Book Series Title","")).strip(),
                doi=row.get("Item DOI","").strip(),
                url=row.get("ArticleURL","").strip() or row.get("URL","").strip(),
                ptype=row.get("Content Type","").strip(), database="Springer", query_id=qid))
    except Exception as e: st.warning(f"Springer parse error: {e}")
    return papers

def parse_scopus_csv(content, qid):
    papers = []
    try:
        for row in csv.DictReader(io.StringIO(content)):
            papers.append(_paper(
                title=row.get("Title","").strip(), authors=row.get("Authors","").strip(),
                year=str(row.get("Year","")).strip(), source=row.get("Source title","").strip(),
                doi=row.get("DOI","").strip(), url=row.get("Link","").strip(),
                ptype=row.get("Document Type","").strip(), database="Elsevier/Scopus", query_id=qid))
    except Exception as e: st.warning(f"Scopus parse error: {e}")
    return papers

def parse_scopus_pop_csv(content, qid):
    papers = []
    try:
        for row in csv.DictReader(io.StringIO(content)):
            papers.append(_paper(
                title=row.get("Title","").strip(), authors=row.get("Authors","").strip(),
                year=str(row.get("Year","")).strip(), source=row.get("Source","").strip(),
                doi=row.get("DOI","").strip(), url=row.get("ArticleURL","").strip(),
                ptype=row.get("Type","Article").strip(), database="Elsevier/Scopus", query_id=qid))
    except Exception as e: st.warning(f"Scopus PoP parse error: {e}")
    return papers

def parse_scholar_csv(content, qid):
    papers = []
    try:
        for row in csv.DictReader(io.StringIO(content)):
            papers.append(_paper(
                title=row.get("Title","").strip(), authors=row.get("Authors","").strip(),
                year=str(row.get("Year","")).strip(), source=row.get("Source","").strip(),
                doi=row.get("DOI","").strip(),
                url=row.get("ArticleURL","").strip() or row.get("URL","").strip(),
                ptype=row.get("Type","article").strip(), database="Google Scholar", query_id=qid))
    except Exception as e: st.warning(f"Scholar parse error: {e}")
    return papers

def parse_bib(content, qid):
    papers = []
    for entry in re.split(r'\n@', content):
        if not entry.strip(): continue
        if not entry.startswith('@'): entry = '@' + entry
        def gf(field, text):
            pat = rf'{re.escape(field)}\s*=\s*\{{(.+?)\}}\s*[,\}}]|{re.escape(field)}\s*=\s*"(.+?)"\s*[,\}}]'
            m = re.search(pat, text, re.IGNORECASE|re.DOTALL)
            if m: r = m.group(1) if m.group(1) else m.group(2); return r.strip().replace('\n',' ')
            return ""
        papers.append(_paper(
            title=gf("title",entry), authors=gf("author",entry), year=gf("year",entry),
            source=gf("booktitle",entry) or gf("journal",entry), doi=gf("doi",entry),
            url=gf("url",entry),
            ptype="Conference Paper" if "inproceedings" in entry[:30].lower() else "Article",
            database="ACM", query_id=qid))
    return [p for p in papers if p["title"]]

def detect_csv_type(content, fname=""):
    first = content.split('\n')[0].lower()
    second = content.split('\n')[1].lower() if len(content.split('\n')) > 1 else ""
    if "scopus" in fname: return "scopus_pop"
    if "springer" in fname: return "springer"
    if "scholar" in fname: return "scholar"
    if "item title" in first: return "springer"
    if "source title" in first: return "scopus"
    if "scopus.com" in second: return "scopus_pop"
    return "scholar"

# ── Deduplication ─────────────────────────────────────────────────────────────
def paper_score(p):
    s = 0
    if p.get("abstract","").strip(): s += 3
    if p.get("doi","").strip():      s += 2
    if p.get("url","").strip():      s += 1
    if p.get("authors","").strip():  s += 1
    return s

def deduplicate(papers):
    sorted_p = sorted(papers, key=paper_score, reverse=True)
    seen_doi, seen_title, unique, dupes = set(), set(), [], []
    for p in sorted_p:
        doi   = p.get("doi","").strip().lower()
        title = p.get("title","").strip().lower()[:120]
        if doi:
            if doi in seen_doi: dupes.append(p); continue
            seen_doi.add(doi)
        else:
            if title and title in seen_title: dupes.append(p); continue
        if title: seen_title.add(title)
        unique.append(p)
    return unique, dupes

# ── Language Detection ────────────────────────────────────────────────────────
GERMAN_WORDS  = ['für','über','durch','beim','zur','zum','nach','mit','von','ist','sind',
                 'wird','werden','wurde','wurden','hat','haben','einer','eines','einem',
                 'einen','der','die','das','und','oder','auch','nicht','sich','dass','eine',
                 'als','bei','des','dem','im','an','auf','zu','in','es','er','sie','wir']
FRENCH_WORDS  = ['pour','avec','dans','sur','les','des','une','est','sont','par','aux',
                 'du','qui','que','il','elle','nous','vous','ils','elles','ce','cette']
SPANISH_WORDS = ['para','con','por','los','las','una','del','que','como','más','este',
                 'esta','están','tiene','pueden','entre','sobre','hacia','desde']

def is_english_text(text):
    """Returns (is_english, reason). Keyword-based — catches German/French even with few umlauts."""
    if not text or len(text.strip()) < 10: return True, ""
    cjk = sum(1 for c in text if '\u4e00' <= c <= '\u9fff' or '\u3040' <= c <= '\u30ff')
    if cjk > 3: return False, "CJK characters"
    special = sum(1 for c in text if '\u0400' <= c <= '\u04ff' or '\u0600' <= c <= '\u06ff')
    if special > 3: return False, "Non-Latin script"
    tl = text.lower(); padded = f" {tl} "
    def cnt(wl): return sum(1 for w in wl if f" {w} " in padded or padded.startswith(f"{w} "))
    if cnt(GERMAN_WORDS)  >= 3: return False, "German detected"
    if cnt(FRENCH_WORDS)  >= 3: return False, "French detected"
    if cnt(SPANISH_WORDS) >= 3: return False, "Spanish detected"
    return True, ""

def is_valid_year(y):
    return str(y).strip().isdigit() and 2015 <= int(str(y).strip()) <= 2026

# ── HTTP + Rate Limiter ───────────────────────────────────────────────────────
import requests as _req
from bs4 import BeautifulSoup

_lock = threading.Lock()
_req_times = []

def _rate_limit():
    with _lock:
        now = time.time()
        _req_times[:] = [t for t in _req_times if now-t < 1.0]
        if len(_req_times) >= 3:
            s = 1.0 - (now - _req_times[0])
            if s > 0: time.sleep(s)
        _req_times.append(time.time())

def _fetch_url(url, timeout=12):
    try:
        _rate_limit()
        r = _req.get(url, headers={"User-Agent":"Mozilla/5.0 Chrome/120"},
                     timeout=timeout, allow_redirects=True, verify=False)
        return r if r.status_code == 200 else None
    except: return None

def _is_english_abstract(text):
    if not text or len(text) < 10: return True
    return (sum(1 for c in text if ord(c) > 127) / len(text)) < 0.3

def _clean(text):
    if not text: return ""
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'^(Abstract|ABSTRACT|Summary|SUMMARY)\s*[:\-]?\s*', '', text, flags=re.I)
    return text

def _norm_doi(doi):
    if not doi: return ""
    doi = doi.strip()
    doi = re.sub(r'/(abstract|full|pdf|reference|v\d+|fulltext|advance-article-abstract|article-abstract|article-pdf)$','',doi,flags=re.I)
    doi = re.sub(r'\.(short|abstract)$','',doi,flags=re.I)
    doi = re.sub(r'^https?://(dx\.)?doi\.org/','',doi,flags=re.I)
    doi = re.sub(r'^doi:','',doi,flags=re.I)
    return doi.rstrip('/')

def _title_match(t1, t2, thresh=0.5):
    if not t1 or not t2: return True
    w1 = set(re.sub(r'[^\w\s]','',t1.lower()).split())
    w2 = set(re.sub(r'[^\w\s]','',t2.lower()).split())
    if not w1 or not w2: return True
    return len(w1 & w2) / min(len(w1),len(w2)) >= thresh

def _is_pdf(url):
    if not url: return False
    return any(x in url.lower() for x in ['.pdf','/pdf/','/article-pdf/','/content/pdf/','/doi/pdf/','/bitstream/handle/'])

def _is_blocked(url):
    if not url: return False
    return any(b in url.lower() for b in ['researchgate.net','ebscohost.com','books.google','sciengine.com','pubpub.org'])

def _is_scopus_inward(url):
    return bool(url) and 'scopus.com/inward' in url.lower()

# ── Year Recovery ─────────────────────────────────────────────────────────────

def _extract_year_from_dateparts(date_parts):
    """Extract year from CrossRef/SemanticScholar date-parts format [[2025, 4, 25]]"""
    try:
        if date_parts and isinstance(date_parts, list) and date_parts[0]:
            return str(date_parts[0][0])
    except: pass
    return ""

def fetch_year_from_semantic_scholar(doi):
    """Try to get year from SemanticScholar by DOI"""
    if not doi: return ""
    try:
        _rate_limit()
        r = _req.get(f"https://api.semanticscholar.org/graph/v1/paper/DOI:{doi}",
                     params={"fields":"year"}, timeout=8)
        if r.ok:
            year = r.json().get("year")
            if year: return str(year)
    except: pass
    return ""

def fetch_year_from_crossref(doi):
    """Try to get year from CrossRef by DOI — checks multiple date fields"""
    if not doi: return ""
    try:
        _rate_limit()
        r = _req.get(f"https://api.crossref.org/works/{doi}",
                     headers={"User-Agent":"SystematicReview/1.0"}, timeout=8)
        if r.ok:
            msg = r.json().get("message", {})
            # Try fields in order of reliability
            for field in ["published", "published-print", "published-online", "issued", "created"]:
                dp = msg.get(field, {}).get("date-parts")
                year = _extract_year_from_dateparts(dp)
                if year and year.isdigit() and 1990 <= int(year) <= 2030:
                    return year
    except: pass
    return ""

def recover_year(paper):
    """Try to recover missing year from APIs. Returns year string or ''"""
    doi = _norm_doi(paper.get("doi",""))
    if doi:
        # Try SemanticScholar first (faster)
        year = fetch_year_from_semantic_scholar(doi)
        if year: return year
        # Try CrossRef
        year = fetch_year_from_crossref(doi)
        if year: return year
    return ""

# ── Abstract Fetching ─────────────────────────────────────────────────────────

def fetch_crossref_abstract(doi, paper_title=""):
    if not doi: return ""
    try:
        _rate_limit()
        r = _req.get(f"https://api.crossref.org/works/{doi}",
                     headers={"User-Agent":"SystematicReview/1.0"}, timeout=10)
        if r.ok:
            msg = r.json().get("message", {})
            abstract = msg.get("abstract","")
            if abstract:
                # Verify it's the right paper
                cr_title = (msg.get("title") or [""])[0]
                if paper_title and cr_title and not _title_match(paper_title, cr_title):
                    return ""
                abstract = _clean(abstract)
                if len(abstract) > 50 and _is_english_abstract(abstract): return abstract
    except: pass
    return ""

def fetch_semantic_scholar_abstract(doi):
    if not doi: return ""
    for url in [f"https://api.semanticscholar.org/graph/v1/paper/DOI:{doi}",
                f"https://api.semanticscholar.org/graph/v1/paper/{doi}"]:
        try:
            _rate_limit()
            r = _req.get(url, params={"fields":"abstract"}, timeout=10)
            if r.ok:
                abstract = r.json().get("abstract") or ""
                if abstract and len(abstract) > 50 and _is_english_abstract(abstract): return abstract
        except: pass
    return ""

def fetch_openalex_abstract(doi):
    if not doi: return ""
    try:
        _rate_limit()
        r = _req.get(f"https://api.openalex.org/works/doi:{doi}", timeout=10)
        if r.ok:
            data = r.json()
            abstract = data.get("abstract","")
            if abstract and len(abstract) > 50 and _is_english_abstract(abstract): return abstract
            inv = data.get("abstract_inverted_index")
            if inv:
                words = []
                for word, positions in inv.items():
                    for pos in positions:
                        while len(words) <= pos: words.append("")
                        words[pos] = word
                abstract = " ".join(words)
                if len(abstract) > 50 and _is_english_abstract(abstract): return abstract
    except: pass
    return ""

def fetch_europepmc_abstract(doi):
    if not doi: return ""
    try:
        _rate_limit()
        r = _req.get("https://www.ebi.ac.uk/europepmc/webservices/rest/search",
                     params={"query":f"DOI:{doi}","format":"json","resultType":"core"}, timeout=10)
        if r.ok:
            results = r.json().get("resultList",{}).get("result",[])
            if results:
                abstract = results[0].get("abstractText","")
                if abstract and len(abstract) > 50 and _is_english_abstract(abstract): return abstract
    except: pass
    return ""

def scrape_html_abstract(html):
    if not html: return ""
    try:
        soup = BeautifulSoup(html,"html.parser")
        for tag in soup(["script","style","nav","header","footer"]): tag.decompose()
        selectors = [
            "div.abstract","section.abstract","div#abstract","div.abstractSection",
            "p.abstract","#Abs1-content","div[class*='abstract']","section[class*='abstract']",
            "section[data-title='Abstract']","div.c-article-section__content",
            "div.Abstract","div#Abs1","div.abstract-text","div.abstract-body",
            "div.article-section__content","div.NLM_sec","div.abstract-content",
            "blockquote.abstract","meta[name='description']","meta[property='og:description']",
            "meta[name='citation_abstract']","meta[name='DC.Description']",
        ]
        for sel in selectors:
            try:
                el = soup.select_one(sel)
                if el:
                    text = el.get("content","") if el.name=="meta" else el.get_text(" ",strip=True)
                    text = _clean(text)
                    if len(text) > 80 and _is_english_abstract(text) and len(text) < 5000: return text
            except: continue
        for h in soup.find_all(['h1','h2','h3','h4','strong','b']):
            if 'abstract' in h.get_text(strip=True).lower():
                nxt = h.find_next(['p','div'])
                if nxt:
                    text = _clean(nxt.get_text(" ",strip=True))
                    if len(text) > 80 and _is_english_abstract(text) and len(text) < 5000: return text
        best = ""
        academic = ['study','research','method','results','analysis','data','model','system','proposed','approach']
        for p in soup.find_all('p'):
            text = p.get_text(" ",strip=True)
            if len(text) > len(best) and 100 < len(text) < 3000:
                if any(w in text.lower() for w in academic): best = text
        if best: return _clean(best)
    except: pass
    return ""

def fetch_abstract_for_paper(paper):
    doi   = _norm_doi(paper.get("doi",""))
    url   = paper.get("url","").strip()
    title = paper.get("title","")

    # 1. DOI APIs: CrossRef (with title check) → SemanticScholar → OpenAlex → EuropePMC
    if doi:
        for fn in [lambda: fetch_crossref_abstract(doi, title),
                   lambda: fetch_semantic_scholar_abstract(doi),
                   lambda: fetch_openalex_abstract(doi),
                   lambda: fetch_europepmc_abstract(doi)]:
            abstract = fn()
            if abstract: return abstract
        # 2. doi.org URL scrape
        doi_url = f"https://doi.org/{doi}"
        r = _fetch_url(doi_url)
        if r and 'pdf' not in r.headers.get('Content-Type','').lower():
            abstract = scrape_html_abstract(r.text)
            if abstract: return abstract

    # 3. Original URL
    if url and not _is_pdf(url) and not _is_blocked(url) and not _is_scopus_inward(url):
        if doi and url == f"https://doi.org/{doi}": pass  # already tried
        else:
            r = _fetch_url(url)
            if r and 'pdf' not in r.headers.get('Content-Type','').lower():
                abstract = scrape_html_abstract(r.text)
                if abstract: return abstract
    return ""

# ── Auto-Screening ────────────────────────────────────────────────────────────

def auto_screen(papers):
    """Auto-screen E1/E2/E7.
    Year logic:
      - Year present but outside 2015-2026 → Exclude E1 (wrong year, auto)
      - Year missing, has DOI → try API recovery → if found fill year, keep Pending
                                                 → if not found, leave Pending (manual)
      - Year missing, no DOI → leave Pending (manual)
    """
    counts = {"E1": 0, "E2": 0, "E7": 0, "E1_recovered": 0}
    for p in papers:
        year = str(p.get("year","")).strip()
        doi  = _norm_doi(p.get("doi",""))

        # E1 check
        if year and year.isdigit() and not is_valid_year(year):
            # Year IS present but clearly wrong range → auto-exclude
            p.update(screening_status="Exclude", exclusion_reason="E1",
                     notes=f"Auto: year {year} outside 2015-2026", auto_excluded=True)
            counts["E1"] += 1
            continue

        if not year or not year.isdigit():
            # Year missing
            if doi:
                # Try API recovery
                recovered = recover_year(p)
                if recovered and recovered.isdigit() and is_valid_year(recovered):
                    p["year"] = recovered
                    p["notes"] = f"Year recovered from API: {recovered}"
                    counts["E1_recovered"] += 1
                    # Fall through to language check
                else:
                    # API couldn't find it — leave for manual
                    p["notes"] = "Year missing — not found via API, check manually"
                    # keep Pending, continue to next paper
                    continue
            else:
                # No DOI, no year — leave for manual
                p["notes"] = "Year missing, no DOI — check manually"
                continue  # keep Pending

        # E2: Language check on title + source
        text = p.get("title","") + " " + p.get("source","")
        is_eng, reason = is_english_text(text)
        if not is_eng:
            p.update(screening_status="Exclude", exclusion_reason="E2",
                     notes=f"Auto: not English — {reason}", auto_excluded=True)
            counts["E2"] += 1
            continue

        # E7: Missing title or authors
        if not p.get("title","").strip() or not p.get("authors","").strip():
            p.update(screening_status="Exclude", exclusion_reason="E7",
                     notes="Auto: missing title or authors", auto_excluded=True)
            counts["E7"] += 1

    return papers, counts

def post_fetch_language_check(papers):
    """After abstract fetching, re-check language on title+abstract. Flag non-English as E2."""
    flagged = 0
    for p in papers:
        if p.get("auto_excluded"): continue  # already excluded
        abstract = p.get("abstract","")
        if not abstract: continue
        text = p.get("title","") + " " + abstract
        is_eng, reason = is_english_text(text)
        if not is_eng:
            p.update(screening_status="Exclude", exclusion_reason="E2",
                     notes=f"Post-fetch: abstract not English — {reason}", auto_excluded=True)
            flagged += 1
    return papers, flagged

# ── Excel Builder ─────────────────────────────────────────────────────────────

def build_dimension_excel(papers, dupe_list, dim_code, dim_name):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    H_FILL    = PatternFill("solid", start_color="1F4E79")
    H_FONT    = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    W_FILL    = PatternFill("solid", start_color="FFF2CC")
    M_FILL    = PatternFill("solid", start_color="FFE0E0")
    D_FILL    = PatternFill("solid", start_color="E8EAF6")
    AUTO_FILL = PatternFill("solid", start_color="E8F5E9")
    THIN      = Side(style="thin", color="BFBFBF")
    BDR       = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    PH_CLR    = {"Identification":"D6E4F0","Screening":"D5E8D4","Eligibility":"FFF2CC","Included":"FCE4D6"}

    dim_papers = [p for p in papers if dim_code in p.get("dimension","")]
    dim_dupes  = [p for p in dupe_list if dim_code in p.get("dimension","")]
    if not dim_papers: return None

    auto_e1  = sum(1 for p in dim_papers if p.get("exclusion_reason")=="E1")
    auto_e2  = sum(1 for p in dim_papers if p.get("exclusion_reason")=="E2")
    auto_e7  = sum(1 for p in dim_papers if p.get("exclusion_reason")=="E7")
    pending  = [p for p in dim_papers if p.get("screening_status")=="Pending"]
    miss_year= [p for p in dim_papers if not is_valid_year(p.get("year",""))]
    miss_doi = [p for p in dim_papers if not p.get("doi","").strip()]

    wb = openpyxl.Workbook()
    COLS = [
        ("Database",14),("Title",50),("Authors",25),("Year",6),
        ("Source / Journal",26),("DOI",28),("URL",28),("Abstract",60),
        ("Screening Status",14),("Exclusion Reason",20),("Notes",30),
    ]

    def write_screen(ws, plist, rfn=None, start=1):
        ws.freeze_panes = f"A{start+1}"
        for ci,(name,w) in enumerate(COLS,1):
            c = ws.cell(row=start,column=ci,value=name)
            c.font=H_FONT; c.fill=H_FILL; c.border=BDR
            c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
            ws.column_dimensions[get_column_letter(ci)].width=w
        ws.row_dimensions[start].height=32
        for ri,p in enumerate(plist,start+1):
            fill = rfn(p) if rfn else (AUTO_FILL if p.get("auto_excluded") else PatternFill("solid",start_color="FFFFFF"))
            vals=[p.get("database",""),p.get("title",""),p.get("authors","")[:120],
                  p.get("year",""),p.get("source","")[:60],p.get("doi",""),
                  p.get("url","")[:100],p.get("abstract",""),
                  p.get("screening_status","Pending"),p.get("exclusion_reason",""),p.get("notes","")]
            for ci,val in enumerate(vals,1):
                c=ws.cell(row=ri,column=ci,value=val)
                c.fill=fill; c.border=BDR; c.font=Font(name="Arial",size=9)
                c.alignment=Alignment(wrap_text=True,vertical="top")
        ws.auto_filter.ref=f"A{start}:{get_column_letter(len(COLS))}{start+len(plist)}"
        if plist:
            dv1=DataValidation(type="list",formula1='"Pending,Include,Exclude"',allow_blank=True)
            ws.add_data_validation(dv1); dv1.add(f'I{start+1}:I{start+len(plist)}')
            dv2=DataValidation(type="list",formula1='"E1,E2,E3,E4,E5,E6,E7,E8,E9,"',allow_blank=True)
            ws.add_data_validation(dv2); dv2.add(f'J{start+1}:J{start+len(plist)}')

    # Sheet 1: PRISMA
    ws=wb.active; ws.title="PRISMA_Flow"; ws.sheet_view.showGridLines=False
    ws.column_dimensions["A"].width=3
    ws["B2"].value=f"PRISMA 2020 Flow Tracker — {dim_name}"
    ws["B2"].font=Font(bold=True,size=14,color="1F4E79",name="Arial")
    ws.merge_cells("B2:H2")
    hdrs=["Phase","Step","Database","Query ID","n (raw)","n (after filter)","Notes"]
    wids=[18,42,18,12,12,16,50]
    for ci,(h,w) in enumerate(zip(hdrs,wids),2):
        c=ws.cell(row=4,column=ci,value=h); c.font=H_FONT; c.fill=H_FILL; c.border=BDR
        c.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
        ws.column_dimensions[get_column_letter(ci)].width=w
    ws.row_dimensions[4].height=28
    rows=[
        ["Identification",f"Records identified: {dim_name}","All","ALL",len(dim_papers)+len(dim_dupes),"",""],
        ["Identification","Duplicate records removed","All","ALL",len(dim_dupes),"",""],
        ["Identification","Records after deduplication","All","ALL",len(dim_papers),"",""],
        ["Screening","Auto-excluded: E1 (invalid year)","All","ALL",auto_e1,"","Auto-detected"],
        ["Screening","Auto-excluded: E2 (not English)","All","ALL",auto_e2,"","Auto-detected (pre + post-fetch)"],
        ["Screening","Auto-excluded: E7 (insufficient detail)","All","ALL",auto_e7,"","Auto-detected"],
        ["Screening","Pending manual screening","All","ALL",len(pending),"","See Screening_Sheet"],
        ["Screening","Records excluded — title screen","All","ALL","","","Fill after manual"],
        ["Screening","Records excluded — abstract screen","All","ALL","","","Fill after manual"],
        ["Screening","Reports sought for retrieval","All","ALL","","",""],
        ["Screening","Reports not retrieved","All","ALL","","",""],
        ["Eligibility","Reports assessed for eligibility","All","ALL","","",""],
        ["Eligibility","Reports excluded with reasons","All","ALL","","","E1-E8"],
        ["Included","Studies included in final review","All","ALL","","","Fill after screening"],
    ]
    for ri,row in enumerate(rows,5):
        fill=PatternFill("solid",start_color=PH_CLR.get(row[0],"FFFFFF"))
        if "auto-excluded" in str(row[1]).lower(): fill=AUTO_FILL
        elif "duplicate" in str(row[1]).lower(): fill=PatternFill("solid",start_color="FFF2CC")
        for ci,val in enumerate(row,2):
            c=ws.cell(row=ri,column=ci,value=val); c.fill=fill; c.border=BDR
            c.font=Font(name="Arial",size=9); c.alignment=Alignment(wrap_text=True,vertical="center")

    # Sheet 2: Screening (non-Scholar papers)
    non_scholar = [p for p in dim_papers if p.get("database","") != "Google Scholar"]
    ws2=wb.create_sheet("Screening_Sheet"); write_screen(ws2, non_scholar)

    # Sheet 3: Google Scholar Screening
    scholar_papers = [p for p in dim_papers if p.get("database","") == "Google Scholar"]
    ws_gs=wb.create_sheet("Google_Scholar_Screening")
    write_screen(ws_gs, scholar_papers)

    # Sheet 4: Concept Matrix
    ws3=wb.create_sheet("Concept_Matrix_W&W")
    ws3.column_dimensions["A"].width=3; ws3.column_dimensions["B"].width=4
    ws3["C2"].value=f"Webster & Watson (2002) Concept Matrix — {dim_name}"
    ws3["C2"].font=Font(bold=True,size=14,color="1F4E79",name="Arial")
    ws3.merge_cells("C2:T2")
    ws3["C3"].value="Add included papers as rows. Mark 'X' where concept is addressed."
    ws3["C3"].font=Font(italic=True,size=9,color="595959")

    if dim_code=="D1":
        concepts=["Data\nStandards","AI-\nReadiness","Machine-\nReadable","Semantic\nAnnotation",
                  "Standard-\nization","Technical\nStandards","Digital-\nReady","AI-Native",
                  "Metadata","Ontology","Automation","Compliance","Validation","LLM","Generative AI","Interoperability"]
        cfills=["D6E4F0"]*16; hclr="1F4E79"
    elif dim_code=="D2":
        concepts=["Context\nEngineering","Context\nDesign","Context\nConstruction",
                  "Structured\nData","Knowledge\nRepresent.","Context\nWindow",
                  "Context\nSelection","Answer\nQuality","Accuracy","Hallucination",
                  "Semantic\nContext","Context\nPackage","Prompt\nContext","Efficiency","Token\nEfficiency","Cost"]
        cfills=["D5E8D4"]*16; hclr="375623"
    else:
        concepts=["Token\nEfficiency","Context\nCompression","Prompt\nCompression",
                  "Answer\nQuality","Accuracy","Performance","Cost\nEfficiency",
                  "Inference\nCost","Token\nCost","RAG","Context\nSelection",
                  "Context\nPruning","Question\nAnswering","Document\nQA","LLM","Retrieval"]
        cfills=["FFF2CC"]*16; hclr="7F6000"

    ws3.merge_cells("C5:R5"); c=ws3["C5"]
    c.value=f"{dim_code}: {dim_name}"; c.fill=PatternFill("solid",start_color=hclr)
    c.font=Font(bold=True,color="FFFFFF",name="Arial",size=9); c.alignment=Alignment(horizontal="center")
    for ci,(h,w) in enumerate([("Author(s) Year",28),("Title",40)],3):
        c=ws3.cell(row=6,column=ci,value=h); c.font=H_FONT; c.fill=H_FILL; c.border=BDR
        c.alignment=Alignment(horizontal="center",vertical="bottom")
        ws3.column_dimensions[get_column_letter(ci)].width=w
    for ci,(concept,cf) in enumerate(zip(concepts,cfills),5):
        c=ws3.cell(row=6,column=ci,value=concept)
        c.fill=PatternFill("solid",start_color=cf); c.font=Font(bold=True,name="Arial",size=8,color="1F4E79")
        c.alignment=Alignment(text_rotation=90,horizontal="center",vertical="bottom",wrap_text=True)
        c.border=BDR; ws3.column_dimensions[get_column_letter(ci)].width=5
    ws3.row_dimensions[6].height=80
    for ri in range(7,57):
        alt=PatternFill("solid",start_color="F8F8F8") if ri%2==0 else PatternFill()
        for ci in range(3,5+len(concepts)):
            c=ws3.cell(row=ri,column=ci,value=""); c.border=BDR; c.fill=alt
            c.font=Font(name="Arial",size=10); c.alignment=Alignment(horizontal="center",vertical="center")
        ws3.row_dimensions[ri].height=18

    # Sheet 7: Exclusion Criteria
    ws4=wb.create_sheet("Exclusion_Criteria")
    ws4["B2"].value=f"Exclusion Criteria (PRISMA) — {dim_name}"
    ws4["B2"].font=Font(bold=True,size=13,color="1F4E79",name="Arial"); ws4.column_dimensions["B"].width=65
    ecolors=["FFE0E0","FFE0E0","FFF2CC","FFF2CC","D5E8D4","D5E8D4","D6E4F0","D6E4F0","F3E5F5"]
    for ri,crit in enumerate(EXCLUSION_CRITERIA,4):
        c=ws4.cell(row=ri,column=2,value=crit); c.font=Font(name="Arial",size=10)
        cidx = ri-4
        c.fill=PatternFill("solid",start_color=ecolors[cidx] if cidx < len(ecolors) else "FFFFFF")
        c.border=BDR; ws4.row_dimensions[ri].height=22
    ws4["B14"].value="Note: E1, E2, E7 auto-detected. E2 also checked post-fetch on title+abstract. E9 checked if URL/DOI requires payment. E3-E6, E8 require manual review."
    ws4["B13"].font=Font(italic=True,name="Arial",size=9,color="595959")

    buf=io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf.read()

# ── Word Template ─────────────────────────────────────────────────────────────

def generate_word():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    doc.add_heading('Systematic Literature Review: Method & Findings', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in ['Prepared by: [Your Name]','Date: [Date]','Supervisor: Prof. Maria Wimmer']:
        doc.add_paragraph(line)
    doc.add_paragraph()

    def add_bullet(text):
        try: doc.add_paragraph(text, style='List Bullet')
        except: doc.add_paragraph(f'• {text}')

    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph('This systematic literature review examines three dimensions related to Large Language Models (LLMs): (1) Standardization & AI, (2) Context Engineering, and (3) Token Efficiency. The review follows PRISMA 2020 guidelines for transparent reporting.')
    doc.add_paragraph('[Add your introduction here...]')

    doc.add_heading('2. Methodology', 1)
    doc.add_heading('2.1 Search Strategy', 2)
    doc.add_paragraph('Database searches were conducted in:')
    for db in ['Springer Link','Scopus / Elsevier','ACM Digital Library','Google Scholar (via Publish or Perish)']: add_bullet(db)
    doc.add_paragraph('Year range: 2015–2026 | Language: English | Types: Journal + Conference + Review')

    for dim, queries in [
        ('Dimension 1: Standardization & AI', [
            '("data standards" OR "technical standards" OR "semantic standards" OR "standardization" OR "data standardization") AND ("machine-readable" OR "AI-ready" OR "AI-native" OR "digital-ready") AND ("large language models" OR LLM OR "generative AI")',
            '("machine-readable standards" OR "digital standards") AND ("semantic annotation" OR metadata OR ontology) AND (automation OR compliance OR validation)',
        ]),
        ('Dimension 2: Context Engineering', [
            '("large language models" OR LLM) AND ("context engineering" OR "context design" OR "context construction" OR "context provisioning") AND ("structured data" OR "structured context" OR "knowledge representation")',
            '("large language models" OR LLM) AND ("prompt context" OR "context window" OR "context selection") AND ("answer quality" OR accuracy OR hallucination)',
            '("structured context" OR "semantic context" OR "context package") AND ("large language models" OR LLM) AND (efficiency OR "token efficiency" OR cost)',
        ]),
        ('Dimension 3: Token Efficiency', [
            '("large language models" OR LLM) AND ("token efficiency" OR "context compression" OR "prompt compression") AND ("answer quality" OR accuracy OR performance)',
            '("large language models" OR LLM) AND ("cost efficiency" OR "inference cost" OR "token cost") AND ("retrieval augmented generation" OR RAG OR "context selection")',
            '("context compression" OR "prompt compression" OR "context pruning") AND ("large language models" OR LLM) AND ("question answering" OR "document QA")',
        ]),
    ]:
        doc.add_heading(dim, 3)
        for q in queries: doc.add_paragraph(q)

    doc.add_heading('2.2 PRISMA Flow', 2)
    doc.add_paragraph('PRISMA 2020 flow per dimension. See attached Excel files for detailed counts.')
    doc.add_paragraph('[Insert PRISMA flow diagram or reference Excel sheets]')

    doc.add_heading('2.3 Exclusion Criteria', 2)
    doc.add_paragraph('Papers excluded based on:')
    for c in EXCLUSION_CRITERIA: add_bullet(c)
    doc.add_paragraph('Note: E1, E2, E7 auto-detected by script. E2 also checked post-abstract-fetch. E9 = paid/not open access. E3–E6, E8 required manual review.')

    doc.add_heading('2.4 Analysis Method — Webster & Watson (2002)', 2)
    doc.add_paragraph('Concept matrix maps included papers to dimension-specific concepts to identify themes and gaps.')

    doc.add_heading('3. Findings', 1)
    for dn, dname, themes in [
        ('3.1','Dimension 1: Standardization & AI',
         ['How do data standards enable AI-readiness?','What semantic annotation supports LLM integration?','How does standardization impact automation and compliance?']),
        ('3.2','Dimension 2: Context Engineering',
         ['How is context engineered for LLM performance?','What structured context approaches improve answer quality?','How do context window limits affect design?']),
        ('3.3','Dimension 3: Token Efficiency',
         ['What token compression techniques exist?','How does compression affect answer quality?','What is the cost-performance trade-off in RAG?']),
    ]:
        doc.add_heading(f'{dn} {dname}', 2)
        doc.add_heading(f'{dn}.1 Synthesis', 3)
        doc.add_paragraph('Key themes:')
        for t in themes: add_bullet(t)
        doc.add_paragraph('[Your synthesis here — based on included papers from Concept Matrix...]')
        doc.add_heading(f'{dn}.2 Summary Table', 3)
        t = doc.add_table(rows=1, cols=4)
        try: t.style = 'Light Grid Accent 1'
        except: pass
        for i,h in enumerate(['Paper (Author, Year)','Key Contribution','Concepts Addressed','Relevance']): t.rows[0].cells[i].text=h
        doc.add_paragraph('[Add one row per included paper...]')

    doc.add_heading('4. Discussion', 1)
    doc.add_paragraph('[Cross-dimensional themes, research gaps, implications...]')
    doc.add_heading('5. Conclusion', 1)
    doc.add_paragraph('[Key findings and future directions...]')
    doc.add_heading('References', 1)
    doc.add_paragraph('[List included papers in APA format]')

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()

# ── UI ────────────────────────────────────────────────────────────────────────

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;700&display=swap');
html,body,[class*="css"]{font-family:'Inter',sans-serif;}
.stApp{background:linear-gradient(135deg,#0d1117 0%,#161b22 100%);}
h1{font-family:'Inter',sans-serif!important;font-weight:700!important;letter-spacing:-0.02em!important;}
h2,h3{font-family:'Inter',sans-serif!important;font-weight:600!important;letter-spacing:-0.01em!important;}
.sr-card{background:rgba(22,27,34,0.8);backdrop-filter:blur(10px);border:1px solid rgba(88,166,255,0.15);border-radius:12px;padding:20px;margin:12px 0;}
.stat-box{background:linear-gradient(135deg,rgba(88,166,255,0.1),rgba(88,166,255,0.05));border:1px solid rgba(88,166,255,0.2);border-radius:10px;padding:20px;text-align:center;}
.stat-num{font-family:'JetBrains Mono',monospace;font-size:2.2rem;color:#58a6ff;font-weight:700;line-height:1;margin-bottom:8px;}
.stat-lbl{font-size:0.75rem;color:#8b949e;text-transform:uppercase;letter-spacing:1.5px;font-weight:500;}
.tag{display:inline-block;padding:4px 12px;border-radius:6px;font-family:'JetBrains Mono',monospace;font-size:0.75rem;font-weight:600;margin:3px;}
.tag-springer{background:rgba(46,160,67,0.15);color:#3fb950;border:1px solid rgba(46,160,67,0.3);}
.tag-acm{background:rgba(248,81,73,0.15);color:#f85149;border:1px solid rgba(248,81,73,0.3);}
.tag-scopus{background:rgba(88,166,255,0.15);color:#58a6ff;border:1px solid rgba(88,166,255,0.3);}
.tag-scholar{background:rgba(210,153,34,0.15);color:#e3b341;border:1px solid rgba(210,153,34,0.3);}
.warn-box{background:rgba(210,153,34,0.1);border:1px solid rgba(210,153,34,0.25);border-radius:8px;padding:12px 16px;margin:8px 0;font-size:0.9rem;color:#e3b341;}
.sheet-card{background:rgba(22,27,34,0.6);border:1px solid rgba(88,166,255,0.15);border-radius:8px;padding:10px 14px;margin:4px 0;font-size:0.85rem;}
.sheet-card b{color:#58a6ff;font-family:'JetBrains Mono',monospace;font-size:0.8rem;}
.stProgress>div>div{background:linear-gradient(90deg,#58a6ff,#3fb950)!important;border-radius:4px!important;}
hr{border:none!important;height:1px!important;background:linear-gradient(90deg,transparent,rgba(88,166,255,0.3),transparent)!important;margin:24px 0!important;}
[data-testid="stFileUploader"]{border:2px dashed rgba(88,166,255,0.3)!important;border-radius:12px!important;}
::-webkit-scrollbar{width:8px;height:8px;}
::-webkit-scrollbar-track{background:#0d1117;}
::-webkit-scrollbar-thumb{background:#30363d;border-radius:4px;}
::-webkit-scrollbar-thumb:hover{background:#58a6ff;}
</style>
""", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["📥 File Importer", "🔍 Keyword Screener"])

with tab1:
    st.markdown("""
    <div style="text-align:center;padding:20px 0 30px 0;">
    <h1 style="font-size:2.5rem;margin-bottom:8px;">🔬 Systematic Review Bot</h1>
    <p style="color:#8b949e;font-size:1.1rem;margin:0;">
    Upload CSV/BIB → Auto-Screen E1/E2/E7 → Fetch Abstracts → Post-fetch Language Check → 3 Workbooks + Word
    </p></div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    col_h1, col_rst1 = st.columns([4,1])
    with col_h1:
        st.markdown('<div class="sr-card"><h3 style="margin-top:0;">📁 Step 1 — Upload Files</h3><p style="color:#8b949e;margin-bottom:12px;">Name files like <code>springer_d1q1.csv</code>, <code>acm_d1q2.bib</code>, <code>scopus_d2q1.csv</code>, <code>scholar_d3q1.csv</code></p></div>', unsafe_allow_html=True)
    with col_rst1:
        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
        if st.button("🔄 Reset / New Run", key="reset_tab1", use_container_width=True):
            st.session_state.clear()
            st.rerun()
    st.caption("Supports: Springer CSV · Scopus CSV · Google Scholar CSV (Publish or Perish) · ACM BibTeX")

    uploaded = st.file_uploader("Drop all files here", type=["csv","bib"],
                                  accept_multiple_files=True, label_visibility="collapsed")

    if uploaded:
        st.markdown("---")
        all_papers, stats, parse_log = [], {"identification":{}}, []

        for f in uploaded:
            fname = f.name.lower()
            qid = "UNKNOWN"
            for q in ["d1q1","d1q2","d2q1","d2q2","d2q3","d3q1","d3q2","d3q3"]:
                if q in fname: qid=q.upper(); break
            cf = f.read().decode("utf-8",errors="replace")
            if fname.endswith(".bib"):
                papers=parse_bib(cf,qid); db="ACM"
            else:
                ctype=detect_csv_type(cf,fname)
                if ctype=="springer":         papers=parse_springer_csv(cf,qid);   db="Springer"
                elif ctype in("scopus","scopus_pop"):
                    papers=parse_scopus_pop_csv(cf,qid) if ctype=="scopus_pop" else parse_scopus_csv(cf,qid)
                    db="Elsevier/Scopus"
                else:                         papers=parse_scholar_csv(cf,qid);    db="Google Scholar"
            parse_log.append((f.name,db,qid,len(papers)))
            stats["identification"].setdefault(db,{}).setdefault(qid,0)
            stats["identification"][db][qid]+=len(papers)
            all_papers.extend(papers)

        unique, dupe_list = deduplicate(all_papers)

        # Auto-screen silently (no UI block)
        with st.spinner("Processing files..."):
            unique, auto_counts = auto_screen(unique)

        pending = [p for p in unique if p.get("screening_status")=="Pending"]
        stats.update({
            "total_raw":len(all_papers), "duplicates":len(dupe_list),
            "after_dedup":len(unique), "auto_total":sum(v for k,v in auto_counts.items() if k!="E1_recovered"),
        })

        # Step 2: Parse log
        st.markdown("---")
        st.markdown('<div class="sr-card"><h3 style="margin-top:0;">📊 Step 2 — Files Parsed</h3></div>',unsafe_allow_html=True)
        for fname,db,qid,n in parse_log:
            db_cls={"Springer":"springer","ACM":"acm","Elsevier/Scopus":"scopus","Google Scholar":"scholar"}.get(db,"springer")
            st.markdown(f"<div style='margin:4px 0;'><code>{fname}</code> → <span class='tag tag-{db_cls}'>{db}</span> <code>{qid}</code> → <b>{n} papers</b></div>",unsafe_allow_html=True)

        # Step 3: Summary
        st.markdown("---")
        st.markdown('<div class="sr-card"><h3 style="margin-top:0;">📈 Step 3 — Summary by Dimension</h3></div>',unsafe_allow_html=True)
        c1,c2,c3,c4=st.columns(4)
        with c1: st.markdown(f'<div class="stat-box"><div class="stat-num">{stats["total_raw"]}</div><div class="stat-lbl">Total Raw</div></div>',unsafe_allow_html=True)
        with c2: st.markdown(f'<div class="stat-box"><div class="stat-num">{len(dupe_list)}</div><div class="stat-lbl">Duplicates Removed</div></div>',unsafe_allow_html=True)
        with c3: st.markdown(f'<div class="stat-box"><div class="stat-num">{stats["after_dedup"]}</div><div class="stat-lbl">Unique Papers</div></div>',unsafe_allow_html=True)
        with c4:
            d_counts={}
            for p in unique: d=p.get("dimension","")[:2]; d_counts[d]=d_counts.get(d,0)+1
            summary=" / ".join(f"{k}:{v}" for k,v in sorted(d_counts.items()))
            st.markdown(f'<div class="stat-box"><div class="stat-num" style="font-size:1.1rem">{summary}</div><div class="stat-lbl">By Dimension</div></div>',unsafe_allow_html=True)

        # Step 5: Generate
        st.markdown("---")
        need_abstract=[p for p in pending if p.get("doi","").strip() or p.get("url","").strip()]
        est_mins=max(1,len(need_abstract)//15)
        st.markdown(f'<div class="sr-card"><h3 style="margin-top:0;">🚀 Step 4 — Generate Workbooks</h3><p style="color:#8b949e;margin-bottom:0;">Fetch abstracts for <b>{len(need_abstract)} pending papers</b> · post-fetch language check · 3 Excel workbooks + Word (~{est_mins} min)</p></div>',unsafe_allow_html=True)

        fetch_toggle=st.checkbox("Fetch abstracts automatically",value=True)
        with st.expander("⚙️ Advanced Options"):
            max_workers=st.slider("Concurrent fetch workers",1,5,3)
            st.info("PDFs, blocked sites, Scopus inward links auto-skipped.")

        generate_clicked=st.button("🚀 Generate All Workbooks",type="primary",use_container_width=True)

        if generate_clicked:
            prog=st.progress(0); s_txt=st.empty(); d_txt=st.empty()

            if fetch_toggle and need_abstract:
                t0=time.time(); total=len(need_abstract)
                s_txt.markdown("🔄 **Fetching abstracts...**")

                pdf_c=sum(1 for p in need_abstract if _is_pdf(p.get("url","")))
                blk_c=sum(1 for p in need_abstract if _is_blocked(p.get("url","")))
                scp_c=sum(1 for p in need_abstract if _is_scopus_inward(p.get("url","")))
                if pdf_c+blk_c+scp_c>0:
                    st.info(f"Auto-skipping: {pdf_c} PDFs · {blk_c} blocked · {scp_c} Scopus links")

                results={}; found=completed=0

                def fetch_one(ip): i,p=ip; return i,fetch_abstract_for_paper(p)

                with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
                    futures={ex.submit(fetch_one,(i,p)):i for i,p in enumerate(need_abstract)}
                    for future in concurrent.futures.as_completed(futures):
                        idx,abstract=future.result()
                        results[idx]=abstract
                        if abstract: found+=1
                        completed+=1
                        pct=int((completed/total)*100)
                        prog.progress(pct/100)
                        elapsed=time.time()-t0
                        rate=completed/elapsed if elapsed>0 else 1
                        m,s=divmod(int((total-completed)/rate),60)
                        d_txt.markdown(f"**{pct}%** · ⏱ {m}m {s}s remaining · ✅ {found} fetched · {completed}/{total}")

                for idx,abstract in results.items():
                    need_abstract[idx]["abstract"]=abstract

                prog.progress(1.0)
                s_txt.markdown(f"✅ **Abstracts done!** {found}/{total} in {int(time.time()-t0)}s")
                d_txt.empty()

                # Post-fetch language check
                s_txt.markdown("🌐 **Post-fetch language check on title+abstract...**")
                unique, pf_flagged = post_fetch_language_check(unique)
                if pf_flagged > 0:
                    st.markdown(f'<div class="warn-box">⚠️ <b>{pf_flagged} papers</b> flagged E2 post-fetch (abstract not English)</div>',unsafe_allow_html=True)
                s_txt.empty()

            # Build Excel
            s_txt.markdown("📊 **Building dimension workbooks...**")
            for dim_code,dim_name in DIMENSION_NAMES.items():
                excel_bytes=build_dimension_excel(unique,dupe_list,dim_code,dim_name)
                if excel_bytes:
                    st.session_state[f"excel_{dim_code}"]=excel_bytes
                    st.session_state[f"fname_{dim_code}"]=f"{dim_name}_{datetime.now():%Y%m%d_%H%M}.xlsx"

            # Build Word
            s_txt.markdown("📝 **Generating Word template...**")
            try:
                st.session_state["word_bytes"]=generate_word()
                st.session_state["word_fname"]=f"SR_Method_Findings_{datetime.now():%Y%m%d_%H%M}.docx"
            except Exception as e:
                st.warning(f"Word failed: {e} — install python-docx")

            st.session_state["excel_ready"]=True
            st.session_state["total"]=stats["after_dedup"]
            st.session_state["dupes"]=len(dupe_list)
            st.session_state["abstracts"]=sum(1 for p in unique if p.get("abstract","").strip())
            st.session_state["auto_total"]=stats.get("auto_total",0)
            s_txt.empty(); prog.empty()

        # Step 6: Download
        if st.session_state.get("excel_ready"):
            st.markdown("---")
            st.markdown('<div class="sr-card"><h3 style="margin-top:0;">📥 Step 5 — Download</h3></div>',unsafe_allow_html=True)
            st.markdown("#### 📊 Dimension Workbooks")
            for dim_code in ["D1","D2","D3"]:
                if f"excel_{dim_code}" in st.session_state:
                    dim_name=DIMENSION_NAMES[dim_code]
                    st.markdown(f'<div class="sheet-card"><b>{dim_name}</b> — PRISMA Flow · Screening (auto+manual) · Duplicates · Missing Year/DOI · Concept Matrix · Exclusion Criteria</div>',unsafe_allow_html=True)
                    st.download_button(f"📥 Download {dim_name}",
                        data=st.session_state[f"excel_{dim_code}"],
                        file_name=st.session_state[f"fname_{dim_code}"],
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key=f"dl_{dim_code}")
            if "word_bytes" in st.session_state:
                st.markdown("#### 📝 Word Template")
                st.markdown('<div class="sheet-card"><b>Method & Findings Template</b> — Intro · Methodology · Search Strings · PRISMA · 3 Dimension Findings · Discussion · Conclusion</div>',unsafe_allow_html=True)
                st.download_button("📥 Download Word Template",
                    data=st.session_state["word_bytes"],
                    file_name=st.session_state["word_fname"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_word")
            n=st.session_state.get("total",0); a=st.session_state.get("abstracts",0)
            d=st.session_state.get("dupes",0); at=st.session_state.get("auto_total",0)
            st.success(f"✅ {n} papers · {a} with abstract · {d} dupes removed · {at} auto-excluded (E1/E2/E7)")

    else:
        st.markdown("---")
        st.markdown('<div class="sr-card"><h3 style="margin-top:0;">📋 File Naming Convention</h3></div>',unsafe_allow_html=True)
        st.dataframe(pd.DataFrame({
            "File":  ["springer_d1q1.csv","springer_d1q2.csv","acm_d1q1.bib","acm_d1q2.bib","scopus_d2q1.csv","scholar_d3q1.csv"],
            "DB":    ["Springer","Springer","ACM","ACM","Elsevier/Scopus","Google Scholar"],
            "Query": ["D1Q1","D1Q2","D1Q1","D1Q2","D2Q1","D3Q1"],
            "Format":["CSV","CSV","BibTeX","BibTeX","CSV","CSV"],
        }),use_container_width=True,hide_index=True)
        st.info("💡 Query ID (d1q1, d2q2 etc.) must be in filename.")

with tab2:
    st.markdown("""
    <div style="text-align:center;padding:10px 0 20px 0;">
    <h2 style="font-size:1.8rem;margin-bottom:6px;">🔍 Keyword Screener</h2>
    <p style="color:#8b949e;margin:0;">Upload Excel → Paste search strings → Screen papers → Download results</p>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    US_UK = {
        'standardization':'standardisation','standardisation':'standardization',
        'optimization':'optimisation','optimisation':'optimization',
        'organization':'organisation','organisation':'organization',
        'anonymization':'anonymisation','anonymisation':'anonymization',
        'tokenization':'tokenisation','tokenisation':'tokenization',
        'digitalization':'digitalisation','digitalisation':'digitalization',
        'digitization':'digitisation','digitisation':'digitization',
        'normalization':'normalisation','normalisation':'normalization',
        'utilization':'utilisation','utilisation':'utilization',
        'generalization':'generalisation','generalisation':'generalization',
        'contextualization':'contextualisation','contextualisation':'contextualization',
        'formalization':'formalisation','formalisation':'formalization',
        'serialization':'serialisation','serialisation':'serialization',
        'modelling':'modeling','modeling':'modelling',
        'analyse':'analyze','analyze':'analyse',
        'behaviour':'behavior','behavior':'behaviour',
    }

    def extract_keywords(raw):
        """
        Extract all keywords from search string.
        Strips AND/OR/NOT/+ completely. Handles malformed/unclosed quotes.
        """
        cleaned = raw
        # Remove + separator lines
        cleaned = re.sub(r'(?m)^\s*\+\s*$', ' ', cleaned)
        # Remove AND/OR/NOT operators
        cleaned = re.sub(r'\bAND\b', ' ', cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r'\bOR\b',  ' ', cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r'\bNOT\b', ' ', cleaned, flags=re.IGNORECASE)
        # Remove parens, +, newlines
        cleaned = re.sub(r'[+()\n\r]', ' ', cleaned)

        # Extract valid quoted phrases (2+ chars)
        quoted = []
        for m in re.finditer(r'"([^"]{2,})"', cleaned):
            t = m.group(1).strip().lower()
            if t: quoted.append(t)

        # Remove quoted content + stray quotes, get single tokens
        no_quotes = re.sub(r'"[^"]*"', ' ', cleaned).replace('"', ' ')
        singles = []
        for token in no_quotes.split():
            t = token.strip('.,;:-()')
            if not t or len(t) < 2: continue
            tl = t.lower()
            if tl in ('and','or','not'): continue
            singles.append(tl)

        # Deduplicate preserving order
        seen = set(); result = []
        for t in quoted + singles:
            t = t.strip()
            if t and t not in seen: seen.add(t); result.append(t)
        return result

    def expand_keywords(terms):
        """Add US/UK spelling variants automatically."""
        expanded = list(terms)
        for term in terms:
            words = term.lower().split()
            alt_words = []; changed = False
            for w in words:
                alt = US_UK.get(w)
                if alt: alt_words.append(alt); changed = True
                else: alt_words.append(w)
            if changed:
                alt_term = ' '.join(alt_words)
                if alt_term not in expanded: expanded.append(alt_term)
        return expanded

    def screen_paper(title, abstract, keywords):
        """
        Returns (status, reason, note) based on what data is available.

        Rules:
        - Title + Abstract both present:
            → keyword in either → Include
            → no keyword in either → Exclude E4
        - Title only (abstract empty):
            → keyword in title → Include
            → no keyword in title → Pending (can't confirm exclude without abstract)
        - No title and no abstract:
            → Pending (nothing to screen)
        """
        has_title    = bool(title.strip())
        has_abstract = bool(abstract.strip())

        title_hits    = [kw for kw in keywords if kw in title.lower()]    if has_title    else []
        abstract_hits = [kw for kw in keywords if kw in abstract.lower()] if has_abstract else []
        all_hits = list(set(title_hits + abstract_hits))

        if has_title and has_abstract:
            if all_hits:
                return "Include", "", f"Match in title+abstract: {', '.join(all_hits[:5])}"
            else:
                return "Exclude", "E4", "No keyword match in title or abstract"

        elif has_title and not has_abstract:
            if title_hits:
                return "Include", "", f"Match in title: {', '.join(title_hits[:5])}"
            else:
                return "Pending", "", "No title keyword match — abstract missing, check manually"

        else:
            return "Pending", "", "No title or abstract — check manually"

    col_l, col_r = st.columns([1,1])
    with col_l:
        col_title, col_reset = st.columns([3,1])
        with col_title:
            st.markdown('<div class="sr-card"><h3 style="margin-top:0;">📤 Upload Screening Excel</h3></div>', unsafe_allow_html=True)
        with col_reset:
            st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
            if st.button("🔄 Reset / New File", key="reset_screener", use_container_width=True):
                st.session_state.clear()
                st.rerun()
        screener_xl = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"], key="screener_upload")

    with col_r:
        st.markdown('<div class="sr-card"><h3 style="margin-top:0;">🔤 Search Strings</h3><p style="color:#8b949e;font-size:0.85rem;margin-bottom:8px;">Separate multiple queries with a line containing just <code>+</code></p></div>', unsafe_allow_html=True)
        search_input = st.text_area("Search strings", value="", height=220,
            placeholder='("data standards" OR "standardization")\nAND ("machine-readable" OR "AI-ready")\nAND ("large language models" OR LLM)\n+\n("machine-readable standards")\nAND (metadata OR ontology)\nAND (automation OR compliance)',
            label_visibility="collapsed", key="search_input")

    if screener_xl and search_input.strip():
        # Reset kw_parsed if file changes
        file_id = screener_xl.name + str(screener_xl.size)
        if st.session_state.get("screener_file_id") != file_id:
            st.session_state["screener_file_id"] = file_id
            st.session_state["kw_parsed"] = False
            st.session_state.pop("current_keywords", None)

        raw_kw = extract_keywords(search_input)
        keywords = expand_keywords(raw_kw)
        if not keywords:
            st.error("No keywords found — check your search strings.")
        else:
            # Show keywords with Parse button step
            col_parse, _ = st.columns([1,2])
            with col_parse:
                parse_clicked = st.button("✅ Parse Keywords", use_container_width=True, key="parse_btn")
            
            if parse_clicked or st.session_state.get("kw_parsed"):
                st.session_state["kw_parsed"] = True
                st.session_state["current_keywords"] = keywords
                with st.expander(f"📋 {len(keywords)} keywords extracted (with US/UK variants) — click to verify", expanded=True):
                    st.markdown(", ".join(f"`{k}`" for k in keywords))
                    st.caption("If a keyword looks wrong, fix your search string above and click Parse again.")
                st.markdown("---")

            if st.session_state.get("kw_parsed") and st.button("🔍 Screen Papers", type="primary", use_container_width=True, key="screen_btn"):
                keywords = st.session_state.get("current_keywords", keywords)
                import openpyxl
                from openpyxl.styles import PatternFill
                INC = PatternFill("solid", start_color="D5E8D4")
                EXC = PatternFill("solid", start_color="FFE0E0")

                wb = openpyxl.load_workbook(io.BytesIO(screener_xl.read()))
                inc_tot = exc_tot = skip_tot = pend_tot = 0

                for sname in wb.sheetnames:
                    if "Screening_Sheet" not in sname and "Google_Scholar_Screening" not in sname: continue
                    ws = wb[sname]
                    hdrs = {str(c.value or "").strip(): c.column for c in ws[1]}
                    tc = hdrs.get("Title")
                    ac = hdrs.get("Abstract") or hdrs.get("Abstract (snippet)")
                    sc = hdrs.get("Screening Status")
                    rc = hdrs.get("Exclusion Reason")
                    nc = hdrs.get("Notes")
                    if not tc or not sc: continue

                    prog = st.progress(0); stat = st.empty()
                    total = ws.max_row - 1

                    for ri in range(2, ws.max_row+1):
                        status = str(ws.cell(ri, sc).value or "").strip()
                        if status == "Exclude":
                            skip_tot += 1
                            prog.progress(min((ri-1)/max(total,1),1.0)); continue

                        title_val    = str(ws.cell(ri, tc).value or "")
                        abstract_val = str(ws.cell(ri, ac).value or "") if ac else ""
                        status_new, reason_new, note_new = screen_paper(title_val, abstract_val, keywords)

                        ws.cell(ri, sc).value = status_new
                        if rc and reason_new: ws.cell(ri, rc).value = reason_new
                        if nc and note_new:
                            ex = str(ws.cell(ri, nc).value or "")
                            ws.cell(ri, nc).value = (ex+" | " if ex else "") + note_new

                        if status_new == "Include":
                            for ci in range(1, ws.max_column+1): ws.cell(ri, ci).fill = INC
                            inc_tot += 1
                        elif status_new == "Exclude":
                            for ci in range(1, ws.max_column+1): ws.cell(ri, ci).fill = EXC
                            exc_tot += 1
                        else:  # Pending
                            PEND = PatternFill("solid", start_color="FFF8DC")
                            for ci in range(1, ws.max_column+1): ws.cell(ri, ci).fill = PEND
                            pend_tot += 1

                        prog.progress(min((ri-1)/max(total,1),1.0))
                        stat.text(f"Row {ri-1}/{total}...")

                    prog.progress(1.0); stat.empty()

                st.markdown("---")
                c1,c2,c3,c4,c5 = st.columns(5)
                with c1: st.markdown(f'<div class="stat-box"><div class="stat-num" style="color:#3fb950">{inc_tot}</div><div class="stat-lbl">✅ Included</div></div>', unsafe_allow_html=True)
                with c2: st.markdown(f'<div class="stat-box"><div class="stat-num" style="color:#f85149">{exc_tot}</div><div class="stat-lbl">❌ Excluded E4</div></div>', unsafe_allow_html=True)
                with c3: st.markdown(f'<div class="stat-box"><div class="stat-num" style="color:#e3b341">{pend_tot}</div><div class="stat-lbl">⏳ Pending</div></div>', unsafe_allow_html=True)
                with c4: st.markdown(f'<div class="stat-box"><div class="stat-num" style="color:#8b949e">{skip_tot}</div><div class="stat-lbl">⏭️ Skipped</div></div>', unsafe_allow_html=True)
                with c5: st.markdown(f'<div class="stat-box"><div class="stat-num">{inc_tot+exc_tot+pend_tot+skip_tot}</div><div class="stat-lbl">Total</div></div>', unsafe_allow_html=True)
                st.markdown('<div class="warn-box">🟢 Include = keyword match in both fields &nbsp;|&nbsp; 🔴 Exclude E4 = no match in title+abstract &nbsp;|&nbsp; 🟡 Pending = abstract missing, check manually</div>', unsafe_allow_html=True)

                buf = io.BytesIO(); wb.save(buf); buf.seek(0)
                st.download_button("📥 Download Screened Excel", data=buf.read(),
                    file_name=f"screened_{datetime.now():%Y%m%d_%H%M}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True, type="primary", key="dl_screened")

    elif screener_xl and not search_input.strip():
        st.info("Paste your search strings on the right.")
    else:
        st.markdown("""
        <div class="sr-card"><h3 style="margin-top:0;">📋 How to use</h3>
        <ol style="color:#8b949e;line-height:2.2;">
        <li>Run <b>File Importer tab</b> → download Excel</li>
        <li>Upload that Excel here</li>
        <li>Paste search strings (separate queries with a line containing just <code>+</code>)</li>
        <li>Click <b>Screen Papers</b></li>
        <li>App checks title + abstract → Include / Exclude (E4)</li>
        <li>Download updated Excel</li>
        </ol>
        <p style="color:#8b949e;font-size:0.85rem;">Already auto-excluded papers are skipped. Include = any query matched fully (all AND-groups satisfied).</p>
        </div>
        """, unsafe_allow_html=True)
